"""Word (.docx) converter — extracts text, formatting, and images to markdown."""

import hashlib
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from src.shared.logger import get_logger
from src.shared.types import ConvertResult, ImageAsset, IngestInput

logger = get_logger("ingest.converters.docx")

DOCX_EXTENSIONS = {".docx"}

# Mapping from docx heading styles to markdown levels
HEADING_MAP = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
    "Heading 5": "##### ",
    "Heading 6": "###### ",
}


def can_handle(input_data: IngestInput) -> bool:
    if input_data.type == "file" and input_data.file_path:
        p = Path(input_data.file_path)
        return p.exists() and p.suffix.lower() in DOCX_EXTENSIONS
    return False


def convert(input_data: IngestInput) -> ConvertResult:
    """Docx -> Markdown + extracted images."""
    docx_path = Path(input_data.file_path)
    doc_name = docx_path.stem

    doc = Document(str(docx_path))

    images = _extract_images(doc, doc_name)
    markdown = _build_markdown(doc, images)
    title = _extract_title(markdown, input_data.original_filename or doc_name)

    logger.info("DOCX converted: %s (%d paragraphs, %d images)", doc_name, len(doc.paragraphs), len(images))
    return ConvertResult(markdown=markdown, title=title, images=images)


def _extract_images(doc: Document, doc_name: str) -> list[ImageAsset]:
    """Extract all embedded images from the docx."""
    images = []
    counter = 1

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                ext = Path(rel.target_part.partname).suffix or ".png"
                filename = f"{doc_name}_img_{counter:03d}{ext}"
                images.append(ImageAsset(filename=filename, data=image_data))
                counter += 1
            except Exception as e:
                logger.warning("Failed to extract image: %s", e)

    return images


def _build_markdown(doc: Document, images: list[ImageAsset]) -> str:
    """Convert docx paragraphs to markdown."""
    lines = []
    image_idx = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        # Check for heading
        prefix = HEADING_MAP.get(style_name, "")

        # Check if paragraph contains an image reference
        has_image = any(
            "graphic" in child.tag or "drawing" in child.tag or "pic:" in child.tag
            for child in para._element.iter()
        )

        if has_image and image_idx < len(images):
            img = images[image_idx]
            lines.append(f"![{img.filename}](./images/{img.filename})")
            lines.append("")
            image_idx += 1

        # Build text with inline formatting
        text = _para_to_markdown(para)
        if not text.strip() and not has_image:
            lines.append("")
            continue

        # List items
        if style_name.startswith("List Bullet") or "Bullet" in style_name:
            lines.append(f"- {text}")
        elif style_name.startswith("List Number") or "Number" in style_name:
            lines.append(f"1. {text}")
        elif prefix:
            lines.append(f"{prefix}{text}")
        else:
            lines.append(text)

        lines.append("")

    return "\n".join(lines).strip()


def _para_to_markdown(para) -> str:
    """Convert a paragraph's runs to markdown with inline formatting."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def _extract_title(markdown: str, fallback: str) -> str:
    """Extract title from first H1 heading, or use filename."""
    for line in markdown.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("##"):
            return stripped[2:].strip()
    return Path(fallback).stem if "." in fallback else fallback
